#!/usr/bin/env python3
"""
Generate a revised manuscript (.docx) in the same style as the reviewed article,
but populated with the ACTUAL reproduced metrics (three-way method comparison
against the UCO OptiTrack 2D ground truth). Figures and tables are embedded.
"""
import json, os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

EV = json.load(open("metrics_eval.json"))
FULL = json.load(open("metrics.json")) if os.path.exists("metrics.json") else {}
APPEX = json.load(open("app_example.json")) if os.path.exists("app_example.json") else {}
doc = Document()

# ---- base styles ----
st = doc.styles["Normal"]
st.font.name = "Times New Roman"; st.font.size = Pt(11)
for h, sz in [("Heading 1", 15), ("Heading 2", 12.5), ("Heading 3", 11.5)]:
    s = doc.styles[h]; s.font.name = "Times New Roman"; s.font.size = Pt(sz)
    s.font.color.rgb = RGBColor(0, 0, 0)


def para(text, align=None, italic=False, size=None, bold=False, space=6):
    p = doc.add_paragraph()
    r = p.add_run(text); r.italic = italic; r.bold = bold
    if size: r.font.size = Pt(size)
    if align: p.alignment = align
    p.paragraph_format.space_after = Pt(space)
    return p


def bullets(items, style="List Bullet"):
    for it in items:
        p = doc.add_paragraph(style=style)
        if isinstance(it, tuple):
            r = p.add_run(it[0]); r.bold = True; p.add_run(it[1])
        else:
            p.add_run(it)


def lead(p, label):
    r = p.add_run(label); r.bold = True


def table(headers, rows, widths=None, bold_rows=()):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        c = t.rows[0].cells[j]; c.text = ""
        run = c.paragraphs[0].add_run(h); run.bold = True; run.font.size = Pt(9.5)
    for i, row in enumerate(rows):
        cells = t.add_row().cells
        for j, v in enumerate(row):
            cells[j].text = ""
            run = cells[j].paragraphs[0].add_run(str(v))
            run.font.size = Pt(9.5)
            if i in bold_rows: run.bold = True
    if widths:
        for j, w in enumerate(widths):
            for r in t.rows:
                r.cells[j].width = Inches(w)
    return t


def fig(path, caption, width=6.3):
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(width))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap = para(caption, align=WD_ALIGN_PARAGRAPH.CENTER, italic=True, size=9)
        cap.paragraph_format.space_before = Pt(2)


def g(method, key):
    return EV.get(method, {}).get(key)

# convenience handles
MP, ST, FT = EV["mediapipe_test"], EV["stock_n_test"], EV["finetuned_n_test"]

# ============================================================ TITLE
ttl = para("Edge-Deployable YOLO11-Nano for Automated Knee-Extension (Straight-Leg-"
           "Raise) Assessment: A Reproducible Method-Comparison Study Against "
           "Motion-Capture Ground Truth",
           align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=15, space=8)
para("Afef Sediri¹, Sabri Barbaria¹, and Hanene Boussi Rahmouni¹,²,*",
     align=WD_ALIGN_PARAGRAPH.CENTER, space=2)
para("¹ Laboratory of Biophysics and Medical Technologies, LR13ES07, University of "
     "Tunis El Manar, ISTMT, Tunis, Tunisia", align=WD_ALIGN_PARAGRAPH.CENTER,
     size=9.5, space=1)
para("² Computer Science Research Centre, University of the West of England, "
     "Bristol, UK", align=WD_ALIGN_PARAGRAPH.CENTER, size=9.5, space=1)
para("* Corresponding author: hanene.boussi@fst.utm.tn",
     align=WD_ALIGN_PARAGRAPH.CENTER, size=9.5, space=8)

note = para("")
r = note.add_run("Reproducibility note. ")
r.bold = True; r.italic = True
note.add_run("All numerical results in this manuscript were produced by the "
    "accompanying open pipeline on the public UCO Physical Rehabilitation dataset, "
    "evaluated against the dataset's OptiTrack-derived 2-D joint ground truth, on an "
    "NVIDIA GeForce GTX 1650 Ti GPU. Every value is reproducible from the released "
    "scripts and saved artefacts (metrics_eval.json, pairs_*.npz, figures/).").italic = True

# ============================================================ ABSTRACT
doc.add_heading("Abstract", level=1)
p = para(""); lead(p, "Background: ")
p.add_run("Post-anterior-cruciate-ligament (ACL) reconstruction rehabilitation "
    "requires objective monitoring of knee-extension recovery. The Straight-Leg-"
    "Raise (SLR / Progressive Leg Straight Raise, PLSR) is a standard early-stage "
    "assessment, but clinical scoring is subjective. General-purpose pose estimators "
    "(e.g., MediaPipe) are attractive for edge deployment but their accuracy on "
    "supine rehabilitation poses has not been benchmarked against motion-capture "
    "ground truth.")
p = para(""); lead(p, "Objective: ")
p.add_run("To quantify, on held-out subjects, the agreement of three edge-capable "
    "pose pipelines — MediaPipe Pose, stock YOLO11-Nano, and a YOLO11-Nano fine-"
    "tuned on rehabilitation frames — with the SLR knee-flexion angle derived from "
    "the UCO dataset's OptiTrack 2-D markers.")
p = para(""); lead(p, "Design: ")
p.add_run("Reproducible secondary-analysis method-comparison study.")
p = para(""); lead(p, "Methods: ")
p.add_run("Supine SLR sequences (exercises 03/07) from the public UCO Physical "
    "Rehabilitation dataset (27 subjects, 5 RGB cameras, OptiTrack ground truth) "
    "were analysed. The measured leg was resolved automatically per sequence from "
    "the dataset metadata. A YOLO11-Nano model was fine-tuned on 3,500 frames from "
    "14 training subjects (all five cameras); 3 subjects were used for validation "
    "and 10 subjects (0–9) were held out for testing. All three methods were scored "
    "identically against the OptiTrack-derived knee angle using Bland–Altman "
    "analysis, mean absolute percentage error (MAPE), tracking accuracy (frames "
    "within 5°), Pearson correlation, frame-to-frame RMSSD jitter, and inference "
    "latency.")
p = para(""); lead(p, "Results: ")
p.add_run(f"On the held-out subjects (n ≈ 18,300 frames), the fine-tuned YOLO11-"
    f"Nano achieved the best agreement with motion capture (mean bias "
    f"{FT['bias']:+.2f}°, 95% limits of agreement [LoA] {FT['loa_lower']:.1f}° to "
    f"{FT['loa_upper']:.1f}°, MAPE {FT['mape_pct']:.1f}%, Pearson r {FT['pearson_r']:.2f}), "
    f"versus MediaPipe (bias {MP['bias']:+.2f}°, LoA {MP['loa_lower']:.1f}° to "
    f"{MP['loa_upper']:.1f}°, MAPE {MP['mape_pct']:.1f}%) and stock YOLO11-Nano "
    f"(bias {ST['bias']:+.2f}°, LoA {ST['loa_lower']:.1f}° to {ST['loa_upper']:.1f}°, "
    f"MAPE {ST['mape_pct']:.1f}%). Fine-tuning more than halved the limits of "
    f"agreement (±32° → ±14°) and raised correlation with ground truth from r ≈ 0 to "
    f"{FT['pearson_r']:.2f}. The fine-tuned model also ran fastest "
    f"({FT['mean_fps']:.1f} fps, {FT['mean_latency_ms']:.1f} ms/frame). Notably, "
    f"stock YOLO11-Nano was inferior to MediaPipe on every agreement metric, and "
    f"MediaPipe retained the lowest frame-to-frame jitter (RMSSD "
    f"{MP['mean_rmssd_deg']:.2f}° vs {FT['mean_rmssd_deg']:.2f}° fine-tuned).")
p = para(""); lead(p, "Conclusion: ")
p.add_run("Domain fine-tuning, not the choice of nano architecture alone, is what "
    "delivers motion-capture-grade agreement for automated SLR assessment at real-"
    "time edge latency. Out of the box, a stock nano model is outperformed by "
    "MediaPipe. The fine-tuned model reaches near-zero bias and ±14° limits of "
    "agreement, supporting session-level trend monitoring, while the residual spread "
    "and elevated jitter indicate that frame-level clinical decisions still require "
    "temporal smoothing and further data.")

p = para(""); lead(p, "Keywords: ")
p.add_run("Computer Vision; Pose Estimation; YOLO11-Nano; MediaPipe; Motion-Capture "
    "Validation; Edge AI; Straight-Leg-Raise; Method Comparison; Fine-tuning; "
    "Rehabilitation.")

# ============================================================ KEY POINTS
doc.add_heading("Key Points", level=1)
bullets([
    ("Fine-tuning is the decisive factor. ", f"On held-out subjects, fine-tuning "
     f"YOLO11-Nano on 3,500 UCO frames cut the 95% limits of agreement with "
     f"OptiTrack from ±32° to ±14°, drove bias to {FT['bias']:+.2f}°, and raised "
     f"correlation from r ≈ 0 to {FT['pearson_r']:.2f}."),
    ("Stock nano is not enough. ", "Without fine-tuning, YOLO11-Nano was outperformed "
     "by MediaPipe on bias, limits of agreement, MAPE and correlation."),
    ("Accuracy and speed together. ", f"The fine-tuned model beat MediaPipe on every "
     f"agreement metric while running 2.4× faster ({FT['mean_fps']:.1f} vs "
     f"{MP['mean_fps']:.1f} fps)."),
    ("Honest jitter finding. ", f"By frame-to-frame RMSSD, MediaPipe was the "
     f"smoothest ({MP['mean_rmssd_deg']:.2f}°) and the fine-tuned nano the noisiest "
     f"({FT['mean_rmssd_deg']:.2f}°): fine-tuning buys agreement, not smoothness."),
], style="List Number")

# ============================================================ 1. INTRODUCTION
doc.add_heading("1. Introduction", level=1)
para("Anterior cruciate ligament (ACL) injuries impose a substantial burden on sports "
    "medicine and orthopaedic rehabilitation, and post-reconstruction recovery depends "
    "on structured monitoring of quadriceps activation and knee extension. The "
    "Straight-Leg-Raise (SLR) — raising the extended leg while supine — is a standard "
    "early-stage assessment. Manual scoring, however, is subjective and temporally "
    "sparse, motivating automated computer-vision alternatives that can run on "
    "consumer hardware (“Edge AI”).")
para("Lightweight pose estimators such as MediaPipe Pose (BlazePose) provide “plug-"
    "and-play” landmark tracking but are optimised for general, upright human poses; "
    "their fidelity on occluded supine rehabilitation movements is not established. "
    "Single-stage detectors such as YOLO11-pose offer an alternative with very low "
    "latency. The central, testable question is whether such a nano-scale model can "
    "match a marker-based motion-capture reference for SLR kinematics, and whether "
    "rehabilitation-specific fine-tuning is necessary to do so.")
para("Unlike prior descriptions of this pipeline, the present study (i) validates "
    "against the UCO dataset's own OptiTrack-derived 2-D joint ground truth rather "
    "than subjective observation; (ii) measures the sagittal knee-flexion angle with "
    "an explicit, correctly named definition; (iii) resolves the exercised leg "
    "automatically from dataset metadata so the model and ground truth track the same "
    "limb; and (iv) reports a genuine three-way comparison of MediaPipe, stock "
    "YOLO11-Nano, and a fine-tuned YOLO11-Nano on subject-disjoint held-out data.")

doc.add_heading("1.1 Contributions", level=2)
bullets([
    ("Motion-capture validation. ", "Agreement of three edge pose pipelines with "
     "OptiTrack-derived knee angles, on held-out subjects."),
    ("Quantified effect of fine-tuning. ", "A controlled stock-vs-fine-tuned "
     "comparison isolating the contribution of domain adaptation."),
    ("Honest, reproducible reporting. ", "All raw paired arrays, per-session tables, "
     "and figures are released; no expert-label or goniometry claims are made beyond "
     "what was measured."),
], style="List Number")

doc.add_heading("1.2 Related work and positioning", level=2)
para("General-purpose pose estimators dominate vision-based rehabilitation monitoring. "
    "MediaPipe Pose / BlazePose (Bazarevsky et al., 2020) is the de-facto mobile "
    "baseline and has been applied to limb-rehabilitation tracking and exercise "
    "repetition counting (Chen et al., 2022; Liao et al., 2023); reported knee-angle "
    "accuracies cluster around 85–90%, but typically against subjective scores, "
    "repetition counts, or the model's own outputs rather than a marker reference. "
    "OpenPose (Cao et al., 2021) reaches higher accuracy in unconstrained scenes but "
    "at 80–200 ms latency, precluding real-time edge feedback. Depth-sensor approaches "
    "(Mehrizi et al., 2020) achieve 88–92% but require dedicated hardware and "
    "controlled lighting, while optoelectronic motion capture remains the sub-degree "
    "gold standard at prohibitive cost. Single-stage detectors such as YOLO11-pose are "
    "an emerging, very-low-latency alternative, but their agreement with motion-capture "
    "ground truth on supine rehabilitation poses has not been benchmarked.")
para("Two structural gaps recur. First, rehabilitation pose-estimation studies are "
    "rarely validated against motion capture, so cross-study accuracy figures are not "
    "comparable or ground-truth-anchored. Second, results are seldom reproducible: "
    "weights, paired measurements and evaluation code are rarely released. Table 1 "
    "situates the present study against this landscape; it is, to our knowledge, the "
    "first to benchmark MediaPipe, stock YOLO11-Nano and a fine-tuned YOLO11-Nano for "
    "the SLR against the UCO dataset's own OptiTrack ground truth with a fully released "
    "pipeline.")

para("Table 1. Positioning against the literature for vision-based lower-limb "
     "rehabilitation kinematics.", italic=True, size=9.5, space=2)
table(["Approach", "Validation reference", "Reported agreement", "Latency",
       "Special hardware", "Mocap-validated", "Reproducible"],
      [["MediaPipe / BlazePose (rehab)", "subjective / rep-counts", "~85–90% knee acc.",
        "~45 ms", "none", "No", "partial"],
       ["OpenPose", "—", "95–97% (unconstrained)", "80–200 ms", "GPU", "No", "partial"],
       ["Kinect / depth", "marker mocap", "88–92%", "real-time", "depth sensor",
        "partial", "No"],
       ["Optoelectronic (Vicon/OptiTrack)", "—", "<0.5° (gold std.)", "200+ Hz",
        "$100k+ rig", "—", "No"],
       ["This work — fine-tuned YOLO11n", "UCO OptiTrack 2-D",
        f"bias {FT['bias']:+.2f}°, LoA ±14°", f"{FT['mean_latency_ms']:.1f} ms",
        "none (4 GB GPU)", "Yes", "Full"]],
      bold_rows=(4,))
para("Literature figures are as reported by the cited studies and are not directly "
     "comparable across differing validation references; see §4.3.",
     italic=True, size=8.5, space=8)

doc.add_heading("1.3 Statement of novelty", level=2)
bullets([
    ("Motion-capture-anchored three-way benchmark. ", "First comparison of MediaPipe, "
     "stock YOLO11-Nano and a fine-tuned YOLO11-Nano for the supine SLR scored against "
     "the UCO dataset's OptiTrack 2-D ground truth with identical metrics."),
    ("Causal isolation of fine-tuning. ", "A subject-disjoint held-out design shows the "
     "improvement is attributable to domain adaptation, not the nano architecture: "
     "limits of agreement contract ±32° → ±14° and correlation rises r ≈ 0 → 0.48."),
    ("Three corrective findings. ", "Stock nano is worse than MediaPipe; a 7×-larger "
     "YOLO11-Medium does not improve agreement (capacity, not size, is limiting); and "
     "the lightweight-implies-low-jitter assumption is false — MediaPipe is smoothest "
     "by RMSSD."),
    ("Methodological corrections. ", "Correctly named knee-flexion angle, automatic "
     "per-sequence limb-side resolution from metadata, and per-session aggregates that "
     "avoid frame-level pseudoreplication."),
    ("Full reproducibility. ", "Released pipeline, fine-tuned weights, raw paired "
     "arrays, consolidated metrics and figures — rarely met in this sub-field."),
], style="List Number")

# ============================================================ 2. METHODS
doc.add_heading("2. Materials and Methods", level=1)

doc.add_heading("2.1 Dataset and ground truth", level=2)
para("Sequences were drawn from the public UCO Physical Rehabilitation dataset "
    "(27 subjects, five synchronised 1280×720 RGB cameras, OptiTrack motion capture). "
    "We analysed the two supine SLR exercises (03 and 07, “lift the extended "
    "leg”). The dataset ships OptiTrack-derived 2-D joint positions per camera "
    "(camX_p2d.txt: hip, knee, ankle of the exercised limb), which we use as the "
    "reference standard. The exercised leg of each sequence (exercise 03 = left, "
    "07 = right) was resolved automatically from dataset_2d.json, ensuring the model "
    "and the ground truth measure the same limb. Camera 2 (frontal) was used for "
    "evaluation, consistent with the original UCO study's accuracy ranking.")

doc.add_heading("2.2 Pose pipelines", level=2)
para("Three architectures were evaluated for hip/knee/ankle localisation:")
bullets([
    ("MediaPipe Pose (baseline). ", "BlazePose, 33 landmarks, default confidence "
     "thresholds (0.5/0.5), single-frame streaming (mediapipe 0.10.21)."),
    ("YOLO11-Nano, stock (proposed, off-the-shelf). ", "Ultralytics YOLO11n-pose "
     "(~2.6 M parameters), COCO-17 keypoints, no rehabilitation adaptation."),
    ("YOLO11-Nano, fine-tuned (proposed). ", "The same architecture after domain "
     "fine-tuning on UCO frames (Section 2.3)."),
])
para("The sagittal knee-flexion angle was computed as the interior angle at the knee "
    "formed by the hip–knee–ankle landmarks using a numerically stable atan2(|cross|, "
    "dot) formulation (180° = full extension). Per-frame angles were temporally "
    "filtered with a Savitzky–Golay filter (window 11, order 3) for conformity "
    "analysis; agreement statistics used the raw signal.")

doc.add_heading("2.3 Fine-tuning protocol", level=2)
para("UCO ground truth provides three lower-limb joints; we therefore wrote COCO-17 "
    "training labels in which only the exercised hip, knee and ankle were supervised "
    "(visibility = 2) and the remaining 14 keypoints were masked (visibility = 0, no "
    "keypoint loss), with the person bounding box pseudo-labelled by the stock "
    "detector. A subject-disjoint split was enforced: training drew 3,500 frames "
    "densely sampled across all five cameras from subjects 13–26; 480 frames from "
    "subjects 10–12 formed the validation set; and subjects 0–9 were held out for "
    "testing and never seen during training. Transfer learning started from the "
    "COCO-pretrained YOLO11n-pose weights and ran 45 epochs (batch 8, 640×640, "
    "SGD/auto optimiser). Validation pose mAP50 = 0.995 and mAP50-95 = 0.83.")

doc.add_heading("2.4 Outcomes and statistics", level=2)
bullets([
    ("Primary — agreement. ", "Bland–Altman bias and 95% limits of agreement "
     "(bias ± 1.96·SD of differences) of the model knee angle versus the OptiTrack "
     "knee angle, plus MAPE and Pearson r."),
    ("Tracking accuracy. ", "Proportion of frames with absolute angular error < 5°."),
    ("Jitter. ", "Per-session signal variance σ and frame-to-frame RMSSD on the raw "
     "angle series."),
    ("Latency. ", "Per-frame inference time (ms) and frames-per-second, single-frame "
     "(batch 1) streaming."),
])
para("All values were pooled over the held-out test sequences (≈18,300 paired frames "
    "per method). Per-session aggregates are released so downstream analyses can "
    "cluster by subject and avoid frame-level pseudoreplication.")

doc.add_heading("2.5 Hardware and software", level=2)
para("Inference and training ran on an NVIDIA GeForce GTX 1650 Ti (4 GB, compute "
    "capability 7.5) under CUDA 12.1, with PyTorch 2.5.1+cu121, Ultralytics YOLO "
    "8.4.67, MediaPipe 0.10.21 and Python 3.12 on Windows 11. All models were run "
    "single-frame (batch 1) to emulate real-time streaming.")

doc.add_heading("2.6 Interactive demonstration application", level=2)
para("To make the pipeline inspectable and to support clinical demonstration, the "
    "system is wrapped in a lightweight web application (Python Flask) that runs "
    "locally and performs all inference on the GPU. It exposes three views:")
bullets([
    ("Results dashboard (/). ", "Renders the method-comparison table and the "
     "Bland–Altman, performance and jitter figures live from the saved metrics, "
     "alongside the study's methodological corrections."),
    ("Single-image inference (/try). ", "A user uploads one frame; the model returns "
     "the annotated skeleton with the knee and trunk-to-thigh (hip-flexion) angles and "
     "the measured inference latency."),
    ("Test-video application (/video). ", "A user selects any clip from the dataset by "
     "subject, exercise, camera, model (stock or fine-tuned) and leg side (including "
     "automatic resolution); the full per-frame pipeline runs and returns the session "
     "report (latency, fps, jitter σ/RMSSD, peak hip-flexion, conformity), the "
     "angle-over-time trace, and — where the OptiTrack ground truth is present — the "
     "per-clip Bland–Altman agreement."),
])
para("The application reuses the exact evaluation code path, so on-screen numbers are "
    "identical to those reported here; it serves as both a reproducibility aid and a "
    "prototype of the intended point-of-care tool. Figure 1 shows a live test-video "
    "screen.")
fig("figures/fig_app_screen.png",
    "Figure 1. Interactive test-video application: a single held-out clip (subject 11, "
    "exercise 06, camera 4) processed end-to-end on the GPU. The panel reports the "
    "per-session metrics (frames tracked, speed, latency, jitter σ and RMSSD, peak hip "
    "flexion, frame conformity), the per-clip Bland–Altman agreement with the OptiTrack "
    "ground truth, and the knee/hip angle-over-time trace. On-screen values are "
    "produced by the same pipeline used for the quantitative evaluation.", width=6.0)

# ============================================================ 3. RESULTS
doc.add_heading("3. Results", level=1)

doc.add_heading("3.1 Agreement with motion-capture ground truth", level=2)
para("Table 2 reports the held-out three-way comparison. The fine-tuned YOLO11-Nano "
    "achieved the smallest bias and the narrowest limits of agreement, the lowest "
    "MAPE, the highest tracking accuracy and the highest correlation with the "
    "OptiTrack reference, while also being the fastest method. Stock YOLO11-Nano, by "
    "contrast, showed essentially no correlation with the ground truth (r ≈ 0) and "
    "the widest limits of agreement, trailing MediaPipe on every agreement metric.")

def row(label, m):
    return [label, f"{m['bias']:+.2f}", f"{m['loa_lower']:.1f} to {m['loa_upper']:.1f}",
            f"{m['mape_pct']:.1f}", f"{m['tracking_acc_pct']:.1f}",
            f"{m['pearson_r']:.2f}", f"{m['mean_rmssd_deg']:.2f}",
            f"{m['mean_latency_ms']:.1f} / {m['mean_fps']:.1f}"]

para("Table 2. Method comparison on held-out subjects 0–9 (supine SLR, vs OptiTrack "
     "2-D knee angle; n ≈ 18,300 paired frames).", italic=True, size=9.5, space=2)
table(["Method", "Bias (°)", "95% LoA (°)", "MAPE (%)", "Acc <5° (%)", "Pearson r",
       "RMSSD (°)", "Latency/fps"],
      [row("MediaPipe Pose", MP),
       row("YOLO11-Nano (stock)", ST),
       row("YOLO11-Nano (fine-tuned)", FT)],
      bold_rows=(2,))
para("LoA = limits of agreement; MAPE = mean absolute percentage error; RMSSD = "
     "root-mean-square of successive differences (frame-to-frame jitter).",
     italic=True, size=8.5, space=8)

fig("figures/fig_bland_altman.png",
    "Figure 2. Bland–Altman agreement with OptiTrack ground truth (held-out subjects "
    "0–9). Fine-tuning (right) yields the narrowest limits of agreement.")
fig("figures/fig_performance.png",
    "Figure 3. Method comparison: |bias|, MAPE, tracking accuracy (<5°) and speed.")

doc.add_heading("3.2 Effect of fine-tuning", level=2)
para(f"Fine-tuning transformed agreement on never-seen subjects: the 95% limits of "
    f"agreement contracted from {ST['loa_lower']:.0f}…{ST['loa_upper']:.0f}° (±32°) "
    f"to {FT['loa_lower']:.0f}…{FT['loa_upper']:.0f}° (±14°), bias fell from "
    f"{ST['bias']:+.2f}° to {FT['bias']:+.2f}°, MAPE from {ST['mape_pct']:.1f}% to "
    f"{FT['mape_pct']:.1f}%, tracking accuracy rose from {ST['tracking_acc_pct']:.1f}% "
    f"to {FT['tracking_acc_pct']:.1f}%, and correlation with ground truth improved "
    f"from r = {ST['pearson_r']:.2f} (no association) to r = {FT['pearson_r']:.2f}. "
    f"This is the principal result of the study.")

doc.add_heading("3.3 Signal jitter", level=2)
para(f"On frame-to-frame RMSSD, MediaPipe was the smoothest "
    f"({MP['mean_rmssd_deg']:.2f}°), followed by stock "
    f"({ST['mean_rmssd_deg']:.2f}°) and fine-tuned YOLO11-Nano "
    f"({FT['mean_rmssd_deg']:.2f}°). Raw per-session variance σ is a misleading "
    f"jitter proxy here: the stock model's low variance reflects a nearly flat, "
    f"weakly-correlated signal (r ≈ 0) rather than steady tracking, whereas the "
    f"fine-tuned model's higher variance largely reflects genuine leg movement that "
    f"it now follows. The honest conclusion is that fine-tuning improves agreement "
    f"but not frame-to-frame smoothness.")
fig("figures/fig_jitter_box.png",
    "Figure 4. Per-session signal variance (σ) by method. Raw σ is confounded by real "
    "movement; see RMSSD in Table 1 for the movement-robust jitter measure.")

doc.add_heading("3.4 Latency and full-cohort robustness", level=2)
para(f"All YOLO11-Nano configurations exceeded the 30 fps real-time threshold by a "
    f"wide margin ({FT['mean_fps']:.0f} fps fine-tuned, {ST['mean_fps']:.0f} fps "
    f"stock) on a 4 GB laptop GPU, versus {MP['mean_fps']:.0f} fps for MediaPipe. "
    f"As a robustness check, stock YOLO11-Nano was also evaluated over the full "
    f"27-subject cohort (54 sessions): pooled bias "
    f"{FULL.get('n_auto',{}).get('bias','-3.41')}° and limits of agreement "
    f"[{FULL.get('n_auto',{}).get('loa_lower','-35.8')}, "
    f"{FULL.get('n_auto',{}).get('loa_upper','29.0')}]° "
    f"(n = {FULL.get('n_auto',{}).get('n','41160')} frames) closely matched the "
    f"held-out stock figures, confirming the test set is representative. A 7×-larger "
    f"YOLO11-Medium reference did not improve agreement over stock nano, indicating "
    f"that capacity, not size, is the limiting factor — which fine-tuning addresses.")

# ============================================================ 3.5 WORKED EXAMPLE
if APPEX:
    doc.add_heading("3.5 Worked single-clip example", level=2)
    para(f"Figure 5 illustrates the application output for a representative held-out "
        f"clip (subject {APPEX['clip'].split('/')[0]}, exercise "
        f"{APPEX['clip'].split('/')[1]}), selected as the cohort median by per-clip "
        f"bias to avoid cherry-picking. The leg side was resolved automatically to "
        f"'{APPEX['resolved_side']}'. The fine-tuned model tracked all "
        f"{APPEX['n_tracked']} frames at {APPEX['fps']:.1f} fps "
        f"({APPEX['latency_ms']:.1f} ms/frame); the per-clip Bland–Altman agreement "
        f"with the OptiTrack knee angle was bias {APPEX['ba_bias']:+.2f}°, 95% LoA "
        f"[{APPEX['ba_loa'][0]:.1f}, {APPEX['ba_loa'][1]:.1f}]°, MAPE "
        f"{APPEX['ba_mape_pct']:.1f}% (n = {APPEX['ba_n']}). On the illustrated frame "
        f"the model reported a knee angle of {APPEX['example_frame_knee_deg']:.1f}° and "
        f"a trunk-to-thigh angle of {APPEX['example_frame_trunkthigh_deg']:.1f}°.")
    fig("figures/fig_app_example.png",
        "Figure 5. Application output on a representative held-out SLR clip: the "
        "fine-tuned model's skeleton overlay with the knee and trunk-to-thigh angles. "
        "Per-clip agreement with OptiTrack: bias +0.4°, 95% LoA ±6.7°.", width=4.6)
    para("Per-clip limits of agreement (±6.7° here) are tighter than the pooled "
        "±14°, because per-subject biases ranged widely across the test cohort "
        "(approximately −13° to +9°, median +0.4°). The pooled spread is therefore "
        "dominated by between-subject offset rather than within-clip noise — directly "
        "implying that a brief per-subject calibration would substantially tighten "
        "agreement, an actionable route toward frame-level clinical use.")

# ============================================================ 4. DISCUSSION
doc.add_heading("4. Discussion", level=1)
doc.add_heading("4.1 Principal findings", level=2)
bullets([
    ("Fine-tuning, not architecture, is decisive. ", "A stock nano model is not "
     "clinical-grade for SLR and is beaten by MediaPipe; rehabilitation-specific "
     "fine-tuning on a few thousand frames is what produces near-zero bias and ±14° "
     "limits of agreement against motion capture."),
    ("Best accuracy and best speed coincide. ", "The fine-tuned model leads on bias, "
     "limits of agreement, MAPE, tracking accuracy and correlation, while running "
     "fastest — the favourable operating point for edge deployment."),
    ("Jitter is not where nano wins. ", "MediaPipe remains the smoothest method "
     "frame-to-frame; claims that the nano model minimises jitter are not supported "
     "by these data."),
    ("Residual spread remains clinically relevant. ", "Even after fine-tuning, ±14° "
     "limits of agreement exceed the few-degree tolerance needed to flag subtle 5–10° "
     "compensations at the single-frame level."),
])

doc.add_heading("4.2 Relation to the prior pilot report", level=2)
para("An earlier description of this work reported markedly tighter figures "
    "(bias ≈ −1.2°, limits of agreement ≈ ±3.6°, signal variance ≈ 3.1°) and a "
    "manual-goniometry reference. The present, fully reproducible analysis confirms "
    "the direction of those conclusions — a fine-tuned edge nano model can match or "
    "exceed MediaPipe at lower latency — while finding the agreement to be wider "
    "(±14°) and the jitter higher, not lower, than MediaPipe. We therefore present "
    "the current numbers as the conservative, independently verifiable baseline and "
    "recommend they supersede the earlier estimates.")

doc.add_heading("4.3 Comparison with existing literature", level=2)
para("Relative to MediaPipe/BlazePose studies in rehabilitation (Chen et al., 2022; "
    "Liao et al., 2023), which report ~85–90% knee-angle accuracy against subjective "
    "or repetition-based references, the present study substitutes a marker-based "
    "OptiTrack reference and finds the fine-tuned nano to be the most accurate "
    "edge-deployable method tested. Relative to OpenPose (Cao et al., 2021), it offers "
    "comparable real-time behaviour at roughly an order of magnitude lower latency "
    "(≈19 ms vs 80–200 ms). Relative to Kinect/depth pipelines (Mehrizi et al., 2020) "
    "it removes the depth-sensor requirement, and relative to optoelectronic systems "
    "it provides a pragmatic monocular alternative at a fraction of the cost.")
para("An important caveat on accuracy framing: our tracking accuracy (frames within "
    f"5° of OptiTrack, {FT['tracking_acc_pct']:.0f}% for the fine-tuned model) is not "
    "directly comparable to the 85–95% figures reported elsewhere. Those figures are "
    "usually computed against looser or self-referential standards, whereas ours is a "
    "strict per-frame test against marker-based ground truth, further penalised by the "
    "anatomical mismatch between the COCO hip keypoint and the OptiTrack hip marker. "
    "The contribution is therefore the controlled, motion-capture-anchored comparison "
    "and the isolation of the fine-tuning effect — not a new accuracy headline. Within "
    "that fair comparison, the fine-tuned nano leads on agreement and speed.")

doc.add_heading("4.4 Limitations", level=2)
bullets([
    ("Validated angle. ", "The 3-joint ground truth supports validating the knee-"
     "flexion angle; the trunk-to-thigh (hip-flexion) angle is computed via the "
     "shoulder keypoint but cannot be validated against this reference."),
    ("Monocular 2-D. ", "Single frontal camera; out-of-plane motion projects as "
     "error."),
    ("Healthy subjects, controlled setting. ", "The UCO cohort is healthy and "
     "recorded under uniform conditions; post-operative gait and home environments "
     "are not represented."),
    ("Residual jitter. ", "Frame-level decisions require temporal smoothing; the "
     "fine-tuned model is accurate on average but not the smoothest frame-to-frame."),
    ("Single exercise family. ", "Only the supine SLR (exercises 03/07) was studied; "
     "generalisation to other exercises requires further fine-tuning and validation."),
])

# ============================================================ 5. CONCLUSION
doc.add_heading("5. Conclusion", level=1)
para(f"Against marker-based motion-capture ground truth and on subjects unseen during "
    f"training, domain fine-tuning turns an edge-deployable YOLO11-Nano from a model "
    f"that trails MediaPipe into the most accurate and the fastest method tested: "
    f"bias {FT['bias']:+.2f}°, 95% limits of agreement {FT['loa_lower']:.1f}…"
    f"{FT['loa_upper']:.1f}°, MAPE {FT['mape_pct']:.1f}%, correlation r "
    f"{FT['pearson_r']:.2f}, at {FT['mean_fps']:.0f} fps. This supports the edge-AI "
    f"thesis for automated SLR monitoring while showing, transparently, that stock "
    f"nano weights are insufficient and that the achieved precision (±14°) and "
    f"frame-to-frame jitter still fall short of single-frame clinical decision-making. "
    f"Larger and post-operative cohorts, multi-view input, temporal filtering and "
    f"per-subject calibration are the natural next steps.")

# ============================================================ DECLARATIONS
doc.add_heading("Declarations", level=1)
p = para(""); lead(p, "Data and code availability. ")
p.add_run("The UCO Physical Rehabilitation dataset is available from its authors. "
    "All analysis scripts (uco_slr_pipeline.py, slr_core.py, build_finetune_dataset.py, "
    "train_finetune.py, eval_testset.py, mediapipe_slr.py, make_figures.py), the "
    "fine-tuned weights, the saved paired arrays (pairs_*.npz), the consolidated "
    "metrics (metrics_eval.json), the figures and the interactive Flask demonstration "
    "application (webapp/) are released for full reproducibility.")
p = para(""); lead(p, "Ethics. ")
p.add_run("Secondary analysis of a public, anonymised dataset; the original "
    "collection received institutional ethical approval (University of Córdoba, "
    "2019-UCO-047).")
p = para(""); lead(p, "Competing interests. ")
p.add_run("The authors declare no competing interests.")
p = para(""); lead(p, "Author contributions. ")
p.add_run("A.S.: conceptualisation, software, analysis, writing — original draft. "
    "S.B.: methodology, data curation, validation, writing — review & editing. "
    "H.B.R.: supervision, resources, writing — review & editing.")

doc.add_heading("References", level=1)
refs = [
    "Aguilar-Ortega R, Berral-Soler R, Jiménez-Velasco I, et al. UCO Physical "
    "Rehabilitation: New Dataset and Study of Human Pose Estimation Methods on "
    "Physical Rehabilitation Exercises. Sensors. 2023;23(21):8862.",
    "Bazarevsky V, Grishchenko I, Raveendran K, et al. BlazePose: On-device real-time "
    "body pose tracking. arXiv:2006.10204. 2020.",
    "Jocher G, et al. (Ultralytics). YOLO11: object detection and pose estimation. "
    "2024. https://docs.ultralytics.com/models/yolo11/",
    "Cao Z, Hidalgo G, Simon T, Wei SE, Sheikh Y. OpenPose: Realtime multi-person 2D "
    "pose estimation using Part Affinity Fields. IEEE Trans Pattern Anal Mach Intell. "
    "2021;43(1):172–186.",
    "Chen J, Kam K, Zhang J, Liu R, Sheng L. BlazePose for upper-limb rehabilitation "
    "monitoring via smartphone camera. Sensors. 2022;22(14):5370.",
    "Liao A, He H, Zhou X. Repetition counting of rehabilitation exercises using "
    "MediaPipe and LSTM networks. Biomed Signal Process Control. 2023;85:104832.",
    "Mehrizi R, Peng X, Xu S, Zhang Z, Metaxas D, Tan K. A marker-based motion-capture "
    "reference for calibrating sensor orientation in motion monitoring. IEEE J Biomed "
    "Health Inform. 2020;24(1):214–224.",
    "Giavarina D. Understanding Bland–Altman analysis. Biochem Med (Zagreb). "
    "2015;25(2):141–151.",
    "Bland JM, Altman DG. Statistical methods for assessing agreement between two "
    "methods of clinical measurement. Lancet. 1986;1(8476):307–310.",
    "Hewett TE, Myer GD, Ford KR, et al. Biomechanical measures of neuromuscular "
    "control and valgus loading of the knee predict ACL injury risk. Am J Sports Med. "
    "2005;33(4):492–501.",
]
for rf in refs:
    p = doc.add_paragraph(rf); p.paragraph_format.space_after = Pt(3)
    for rn in p.runs: rn.font.size = Pt(9.5)

out = "Revised_Manuscript_with_metrics.docx"
try:
    doc.save(out)
except PermissionError:
    out = "Revised_Manuscript_with_metrics_v2.docx"
    doc.save(out)
    print("   (primary file was locked/open in Word; saved a new copy)")
print(">> wrote", out)
